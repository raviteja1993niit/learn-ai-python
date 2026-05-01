"""
convert.py — Helper script: converts PDF / DOCX / XLSX to a structured JSON tree.

Supports both documents WITH a Table of Contents (TOC) and WITHOUT:
  - With TOC    : Uses the document's own outline/headings for section titles.
  - Without TOC : Heuristically detects heading-like lines; for truly flat docs,
                  auto-generates a descriptive heading from the first sentence of
                  each section so the JSON tree is always navigable.

Usage:
    python convert.py <input_file> [output.json]

If output path is omitted, writes <input_file_stem>.json next to the input file (same directory).

JSON schema produced:
{
  "filename": "...",
  "file_type": "pdf|docx|xlsx",
  "has_toc": true|false,
  "sections": [
    {
      "id": 0,
      "heading": "Section Title",   // always present — generated if not in doc
      "level": 1,                   // heading depth (1 = top-level)
      "content": "...",             // plain text body of this section
      "page": 1                     // page number (PDF only)
    }
  ],
  "tables": [
    {
      "id": 0,
      "section_id": 0,
      "headers": ["col1", "col2"],
      "rows": [["a", "b"], ["c", "d"]]
    }
  ]
}
"""

import sys
import json
import re
from pathlib import Path


# ── Heading utilities ──────────────────────────────────────────────────────────

def _is_title_case(text: str) -> bool:
    words = text.split()
    if not words:
        return False
    cap = sum(1 for w in words if w and w[0].isupper())
    return cap / len(words) >= 0.7 and len(words) <= 12


def _looks_like_heading(line: str) -> bool:
    """Return True if a line looks like a section heading (not body prose)."""
    line = line.strip()
    if not line or len(line) > 120:
        return False
    # Numbered heading: "1.", "1.2", "Chapter 3", "Section IV"
    if re.match(r'^(chapter|section|part|appendix)\s', line, re.IGNORECASE):
        return True
    if re.match(r'^\d+(\.\d+)*\.?\s+\S', line):
        return True
    # ALL-CAPS short line
    if line.isupper() and 3 <= len(line) <= 80:
        return True
    # Title Case short line (no sentence-ending punctuation)
    if _is_title_case(line) and not line.endswith(('.', ',', ';', ':')):
        return True
    return False


def _auto_heading(content: str, page: int | None = None) -> str:
    """
    Generate a concise heading from the first meaningful sentence of content.
    Falls back to 'Page N' or 'Section' if nothing useful is found.
    """
    if not content:
        return f"Page {page}" if page else "Section"

    # Take first non-trivial line
    for line in content.splitlines():
        line = line.strip()
        if len(line) < 10:
            continue
        # Truncate to first sentence or ~60 chars
        sentence = re.split(r'[.!?]', line)[0].strip()
        if len(sentence) > 8:
            return sentence[:80] + ("…" if len(sentence) > 80 else "")

    return f"Page {page}" if page else "Section"


# ── PDF ────────────────────────────────────────────────────────────────────────

PDF_BATCH_SIZE = 5  # number of pages loaded per batch


def _load_pdf_pages_in_batches(reader, batch_size: int = PDF_BATCH_SIZE) -> list[str]:
    """Load PDF page texts in batches of `batch_size`, returning a flat list."""
    total_pages = len(reader.pages)
    page_texts = []
    for batch_start in range(0, total_pages, batch_size):
        batch_end = min(batch_start + batch_size, total_pages)
        for i in range(batch_start, batch_end):
            page_texts.append(reader.pages[i].extract_text() or "")
    return page_texts


def _pdf_outline_to_flat(outline, reader) -> list[dict]:
    """Recursively flatten a PyPDF2 outline into [{title, page_num, level}]."""
    result = []

    def _walk(items, level=1):
        for item in items:
            if isinstance(item, list):
                _walk(item, level + 1)
            else:
                try:
                    page_num = reader.get_destination_page_number(item) + 1
                except Exception:
                    page_num = None
                result.append({
                    "title": item.title.strip() if hasattr(item, "title") else str(item),
                    "page_num": page_num,
                    "level": level,
                })

    _walk(outline)
    return result


def _parse_pdf(file_bytes: bytes) -> dict:
    import PyPDF2
    import io

    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)

    # ── Extract all page texts in batches of PDF_BATCH_SIZE ───────────────────
    page_texts = _load_pdf_pages_in_batches(reader, PDF_BATCH_SIZE)

    # ── Detect TOC via PDF outline (bookmarks) ────────────────────────────────
    outline = reader.outline if hasattr(reader, "outline") else []
    flat_outline = _pdf_outline_to_flat(outline, reader) if outline else []
    has_toc = len(flat_outline) >= 2

    sections = []
    tables = []

    if has_toc:
        # ── With TOC: use outline entries as section boundaries ───────────────
        # Add a sentinel at the end
        flat_outline.append({"title": None, "page_num": total_pages + 1, "level": 1})

        for i, entry in enumerate(flat_outline[:-1]):
            next_entry = flat_outline[i + 1]
            start = (entry["page_num"] or 1) - 1       # 0-based
            end = (next_entry["page_num"] or total_pages + 1) - 1

            # Collect content from pages in this section
            content_parts = []
            for p in range(max(0, start), min(total_pages, end)):
                content_parts.append(page_texts[p])

            sections.append({
                "id": len(sections),
                "heading": entry["title"],
                "level": entry["level"],
                "content": "\n".join(content_parts).strip(),
                "page": entry["page_num"],
            })

    else:
        # ── Without TOC: scan line-by-line for heading-like lines ─────────────
        current_heading = None
        current_page = 1
        content_lines = []

        def _flush_section():
            nonlocal current_heading, current_page
            body = " ".join(content_lines).strip()
            if not body:
                content_lines.clear()
                return
            heading = current_heading or _auto_heading(body, current_page)
            sections.append({
                "id": len(sections),
                "heading": heading,
                "level": 1,
                "content": body,
                "page": current_page,
            })
            content_lines.clear()

        for page_num, raw in enumerate(page_texts, start=1):
            lines = [l.strip() for l in raw.splitlines() if l.strip()]
            for line in lines:
                if _looks_like_heading(line):
                    _flush_section()
                    current_heading = line
                    current_page = page_num
                else:
                    content_lines.append(line)

        _flush_section()

        # If nothing was split (entire doc is flat body), group pages in batches of PDF_BATCH_SIZE
        if not sections:
            for batch_start in range(0, total_pages, PDF_BATCH_SIZE):
                batch_end = min(batch_start + PDF_BATCH_SIZE, total_pages)
                batch_pages = page_texts[batch_start:batch_end]
                body = "\n".join(t.strip() for t in batch_pages if t.strip())
                if body:
                    first_page = batch_start + 1
                    sections.append({
                        "id": len(sections),
                        "heading": _auto_heading(body, first_page),
                        "level": 1,
                        "content": body,
                        "page": first_page,
                    })

    return {"sections": sections, "tables": tables, "has_toc": has_toc}


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes) -> dict:
    import docx
    import io

    doc = docx.Document(io.BytesIO(file_bytes))
    sections = []
    tables = []

    # Detect whether document uses Heading styles
    heading_styles = {p.style.name for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")}
    has_toc = len(heading_styles) > 0

    current_heading = None
    current_level = 1
    content_lines = []

    def _flush_section():
        nonlocal current_heading, current_level
        body = "\n".join(content_lines).strip()
        if not body and current_heading is None:
            content_lines.clear()
            return
        heading = current_heading or _auto_heading(body)
        sections.append({
            "id": len(sections),
            "heading": heading,
            "level": current_level,
            "content": body,
            "page": None,
        })
        content_lines.clear()

    for block in doc.paragraphs:
        style = block.style.name if block.style else ""
        text = block.text.strip()
        if not text:
            continue

        if style.startswith("Heading"):
            _flush_section()
            m = re.search(r"(\d+)", style)
            current_level = int(m.group(1)) if m else 1
            current_heading = text
        elif has_toc:
            # With TOC: accumulate content under current heading
            content_lines.append(text)
        else:
            # Without TOC: use line-level heuristic heading detection
            if _looks_like_heading(text):
                _flush_section()
                current_heading = text
                current_level = 1
            else:
                content_lines.append(text)

    _flush_section()

    # If still flat (no headings found at all), chunk into groups of ~10 paragraphs
    if not sections:
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        chunk_size = 10
        for i in range(0, len(all_paras), chunk_size):
            chunk = all_paras[i:i + chunk_size]
            body = "\n".join(chunk)
            sections.append({
                "id": len(sections),
                "heading": _auto_heading(body),
                "level": 1,
                "content": body,
                "page": None,
            })

    # Extract tables and associate with nearest preceding section
    for tbl in doc.tables:
        rows_data = []
        headers = []
        for i, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if i == 0:
                headers = cells
            else:
                rows_data.append(cells)
        tables.append({
            "id": len(tables),
            "section_id": max(0, len(sections) - 1),
            "headers": headers,
            "rows": rows_data,
        })

    return {"sections": sections, "tables": tables, "has_toc": has_toc}


# ── XLSX / XLS ─────────────────────────────────────────────────────────────────

def _parse_xlsx(file_bytes: bytes) -> dict:
    import openpyxl
    import io

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sections = []
    tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            continue

        non_empty_rows = [
            [str(c) if c is not None else "" for c in row]
            for row in all_rows
            if any(c is not None for c in row)
        ]
        if not non_empty_rows:
            continue

        headers = non_empty_rows[0]
        data_rows = non_empty_rows[1:]

        content_parts = [
            f"Sheet: {sheet_name}",
            "Columns: " + ", ".join(str(h) for h in headers if h),
            f"Data rows: {len(data_rows)}",
        ]

        sec_id = len(sections)
        sections.append({
            "id": sec_id,
            "heading": sheet_name,
            "level": 1,
            "content": "\n".join(content_parts),
            "page": None,
        })
        tables.append({
            "id": len(tables),
            "section_id": sec_id,
            "headers": headers,
            "rows": data_rows[:500],
        })

    return {"sections": sections, "tables": tables, "has_toc": True}


# ── Main ───────────────────────────────────────────────────────────────────────

def convert(input_path: str, output_path: str | None = None) -> str:
    """Convert a document file to a structured JSON tree. Returns output path."""
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {input_path}")

    file_bytes = path.read_bytes()
    ext = path.suffix.lower()

    if ext == ".pdf":
        data = _parse_pdf(file_bytes)
        file_type = "pdf"
    elif ext in (".docx", ".doc"):
        data = _parse_docx(file_bytes)
        file_type = "docx"
    elif ext in (".xlsx", ".xls"):
        data = _parse_xlsx(file_bytes)
        file_type = "xlsx"
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .doc, .xlsx, .xls")

    result = {
        "filename": path.name,
        "file_type": file_type,
        "has_toc": data.get("has_toc", False),
        "sections": data["sections"],
        "tables": data["tables"],
    }

    out = output_path or str(path.with_suffix(".json"))
    Path(out).write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")

    toc_label = "with TOC" if result["has_toc"] else "without TOC (headings auto-generated)"
    print(f"✓ Converted '{path.name}' ({toc_label}) → '{out}'")
    print(f"  Sections: {len(result['sections'])}, Tables: {len(result['tables'])}")
    return out


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert.py <input_file> [output.json]")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
