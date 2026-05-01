"""
parser.py — Extract per-page / per-section content from documents.

Returns List[Page] — natural document units — no chunking required.
Each Page is a self-contained unit sent directly to the LLM.
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List


@dataclass
class Page:
    page_num: int        # 1-based index
    title: str           # heading or "Page N"
    content: str         # full text of this unit
    word_count: int
    source_type: str     # "pdf_page" | "docx_section" | "xlsx_sheet" | "text_section" | "csv_batch"

    def __post_init__(self):
        if not self.word_count:
            self.word_count = len(self.content.split())


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> List[Page]:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, pdf_page in enumerate(pdf.pages, 1):
                text = pdf_page.extract_text() or ""
                # Also extract tables as pipe-separated rows
                for table in pdf_page.extract_tables() or []:
                    for row in table:
                        row_text = " | ".join(str(c) for c in row if c)
                        if row_text:
                            text += "\n" + row_text
                text = text.strip()
                if text:
                    # Use first non-empty line as title
                    first_line = text.splitlines()[0][:80].strip()
                    title = first_line if first_line else f"Page {i}"
                    pages.append(Page(i, title, text, 0, "pdf_page"))
        return pages
    except ImportError:
        # Fallback to PyPDF2
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, pdf_page in enumerate(reader.pages, 1):
            text = (pdf_page.extract_text() or "").strip()
            if text:
                title = text.splitlines()[0][:80].strip() or f"Page {i}"
                pages.append(Page(i, title, text, 0, "pdf_page"))
        return pages


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes) -> List[Page]:
    import docx as python_docx

    doc = python_docx.Document(io.BytesIO(file_bytes))
    sections: List[tuple[str, List[str]]] = []  # (heading, [paragraphs])
    current_heading = "Introduction"
    current_paras: List[str] = []

    HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "title"}

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if any(h in style_name for h in HEADING_STYLES):
            if current_paras:
                sections.append((current_heading, current_paras))
            current_heading = text[:100]
            current_paras = []
        else:
            current_paras.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                current_paras.append(row_text)

    if current_paras:
        sections.append((current_heading, current_paras))

    # If no heading structure, fall back to paragraph batches
    if not sections:
        all_text = "\n\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        return _split_text_sections(all_text, source_type="docx_section")

    pages = []
    for i, (heading, paras) in enumerate(sections, 1):
        content = "\n\n".join(paras)
        pages.append(Page(i, heading, content, 0, "docx_section"))
    return pages


# ── XLSX ──────────────────────────────────────────────────────────────────────

def _parse_xlsx(file_bytes: bytes) -> List[Page]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    pages = []
    for i, sheet_name in enumerate(wb.sheetnames, 1):
        ws = wb[sheet_name]
        rows = []
        headers: List[str] = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if not any(cells):
                continue
            if row_idx == 0:
                headers = cells
                rows.append("| " + " | ".join(cells) + " |")
                rows.append("|" + "|".join("---" for _ in cells) + "|")
            else:
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            content = "\n".join(rows)
            pages.append(Page(i, f"Sheet: {sheet_name}", content, 0, "xlsx_sheet"))
    return pages


# ── CSV ───────────────────────────────────────────────────────────────────────

def _parse_csv(file_bytes: bytes, rows_per_page: int = 50) -> List[Page]:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    all_rows = list(reader)
    if not all_rows:
        return []

    headers = all_rows[0]
    data_rows = all_rows[1:]
    pages = []

    for batch_idx, start in enumerate(range(0, max(len(data_rows), 1), rows_per_page)):
        batch = data_rows[start:start + rows_per_page]
        lines = ["| " + " | ".join(headers) + " |",
                 "|" + "|".join("---" for _ in headers) + "|"]
        for row in batch:
            lines.append("| " + " | ".join(row) + " |")
        page_num = batch_idx + 1
        title = f"Rows {start + 1}–{min(start + rows_per_page, len(data_rows))}"
        pages.append(Page(page_num, title, "\n".join(lines), 0, "csv_batch"))

    return pages


# ── TXT / Markdown ────────────────────────────────────────────────────────────

def _split_text_sections(text: str, source_type: str = "text_section") -> List[Page]:
    """Split plain text / markdown by headings or double-newline paragraphs."""
    # Try markdown headings first
    heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)
    heading_positions = [(m.start(), m.group(2).strip()) for m in heading_pattern.finditer(text)]

    if len(heading_positions) >= 2:
        sections = []
        for i, (pos, heading) in enumerate(heading_positions):
            end = heading_positions[i + 1][0] if i + 1 < len(heading_positions) else len(text)
            content = text[pos:end].strip()
            # Remove the heading line from content body
            content_body = "\n".join(content.splitlines()[1:]).strip()
            if content_body:
                sections.append((heading, content_body))
        if sections:
            return [
                Page(i + 1, h, c, 0, source_type)
                for i, (h, c) in enumerate(sections)
            ]

    # Fall back: split on double newlines, batch ~800 words per page
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    pages: List[Page] = []
    current: List[str] = []
    current_words = 0
    batch_size = 800

    for para in paragraphs:
        wc = len(para.split())
        if current_words + wc > batch_size and current:
            body = "\n\n".join(current)
            title = current[0][:80].strip()
            pages.append(Page(len(pages) + 1, title, body, 0, source_type))
            current, current_words = [para], wc
        else:
            current.append(para)
            current_words += wc

    if current:
        body = "\n\n".join(current)
        title = current[0][:80].strip()
        pages.append(Page(len(pages) + 1, title, body, 0, source_type))

    return pages if pages else [Page(1, "Document", text[:2000], 0, source_type)]


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_file(filename: str, file_bytes: bytes) -> List[Page]:
    ext = Path(filename).suffix.lower()
    dispatch = {
        ".pdf":  _parse_pdf,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".xls":  _parse_xlsx,
        ".csv":  _parse_csv,
        ".txt":  lambda b: _split_text_sections(b.decode("utf-8", errors="replace"), "text_section"),
        ".md":   lambda b: _split_text_sections(b.decode("utf-8", errors="replace"), "md_section"),
    }
    parser_fn = dispatch.get(ext)
    if not parser_fn:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(dispatch)}")
    pages = parser_fn(file_bytes)
    if not pages:
        raise ValueError("Could not extract any content from the document.")
    return pages

